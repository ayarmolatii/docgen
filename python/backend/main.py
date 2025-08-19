from fastapi import FastAPI
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import uuid
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

app = FastAPI()

# Allow React frontend requests
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # restrict to frontend IP if needed
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/generate-docx")
def generate_docx(req: dict):
    file_path = f"temp_{uuid.uuid4()}.docx"
    doc = Document()

    for comp in req.get("components", []):
        if comp["type"] == "text":
            doc.add_paragraph(comp.get("content", ""))
        elif comp["type"] == "table":
            rows = comp.get("rows", [])
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                for i, row in enumerate(rows):
                    for j, val in enumerate(row):
                        table.cell(i, j).text = str(val)

    doc.save(file_path)
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="generated.docx",
    )

@app.post("/generate-pdf")
def generate_pdf(req: dict):
    file_path = f"temp_{uuid.uuid4()}.pdf"
    c = canvas.Canvas(file_path, pagesize=A4)
    width, height = A4
    y = height - 50  # start from top

    for comp in req.get("components", []):
        if comp["type"] == "text":
            text = comp.get("content", "")
            c.setFont("Helvetica", 12)
            c.drawString(50, y, text)
            y -= 20
            if y < 50:
                c.showPage()
                y = height - 50
        elif comp["type"] == "table":
            rows = comp.get("rows", [])
            if rows:
                col_width = (width - 100) / len(rows[0])
                row_height = 20
                for r, row in enumerate(rows):
                    for c_idx, cell in enumerate(row):
                        x = 50 + c_idx * col_width
                        c.rect(x, y - row_height, col_width, row_height, stroke=1, fill=0)
                        c.drawString(x + 2, y - row_height + 5, str(cell))
                    y -= row_height
                    if y < 50:
                        c.showPage()
                        y = height - 50
                y -= 10  # space after table

    c.save()
    return FileResponse(
        file_path,
        media_type="application/pdf",
        filename="generated.pdf",
    )
